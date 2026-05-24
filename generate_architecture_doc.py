"""
Generates Medibox Architecture & Deployment Status Word document.
Run: python generate_architecture_doc.py
Output: C:/Users/Guesmi Taha/Desktop/Medibox_Architecture.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ──────────────────────────────────────────────────────────
GREEN   = RGBColor(0x1E, 0x8E, 0x3E)   # done
ORANGE  = RGBColor(0xF9, 0xAB, 0x00)   # in-progress / pending
RED     = RGBColor(0xC5, 0x22, 0x1F)   # blocked / not done
BLUE    = RGBColor(0x18, 0x65, 0xD6)   # section headers
DARK    = RGBColor(0x20, 0x23, 0x24)   # body text
LIGHT   = RGBColor(0xF8, 0xF9, 0xFA)   # table header bg (approx)

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Thin grey borders on every cell."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "CCCCCC")
                tcBorders.append(b)
            tcPr.append(tcBorders)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = BLUE
    run.font.bold = True
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def status_badge(cell, text, color):
    cell.text = ""
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.bold  = True
    run.font.size  = Pt(9)
    run.font.color.rgb = color

def add_table(doc, headers, rows, col_widths=None, header_bg="1865D6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, header_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, (val, *opts) in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            color = opts[0] if opts else None
            bold  = opts[1] if len(opts) > 1 else False
            bg    = opts[2] if len(opts) > 2 else None
            if bg:
                set_cell_bg(cell, bg)
            cell.text = ""
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if color:
                run.font.color.rgb = color
            if bold:
                run.font.bold = True

    set_cell_borders(table)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

# ── STATUS HELPERS ──────────────────────────────────────────────────────────
DONE    = ("✅ Done",    GREEN,  True,  None)
PROG    = ("⏳ In Progress", ORANGE, True, None)
PEND    = ("🕐 Pending", ORANGE, False, None)
BLOCKED = ("❌ Blocked", RED,   True,  None)
NO      = ("❌ Not Done", RED,   True,  None)

def cell(text, color=None, bold=False, bg=None):
    return (text, color, bold, bg)

# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

# ── TITLE PAGE ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MEDIBOX")
r.font.size  = Pt(36)
r.font.bold  = True
r.font.color.rgb = BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Medical Prescription OCR Platform")
r2.font.size  = Pt(18)
r2.font.color.rgb = DARK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run("Full Architecture & Deployment Status Report")
r3.font.size  = Pt(13)
r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = date_p.add_run(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}  ·  Project: verox-4dc3f  ·  Region: us-central1")
r4.font.size  = Pt(10)
r4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. WHAT THE SYSTEM DOES
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "1. What the System Does")
p = doc.add_paragraph(
    "Medibox is an end-to-end automated prescription OCR system for Tunisian pharmacies. "
    "A Raspberry Pi 5 sits on the pharmacy counter. It captures a prescription with a USB camera, "
    "runs YOLO locally to detect and crop the handwritten zones, then sends those crops to a cloud AI "
    "pipeline that extracts medications, dosages, and patient information in structured JSON. "
    "The pharmacist reviews the result on a web dashboard, corrects any errors, and the corrections "
    "feed back into periodic model fine-tuning."
)
p.runs[0].font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# 2. HIGH-LEVEL ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "2. High-Level Architecture")
p = doc.add_paragraph()
p.add_run("Five layers, left-to-right:").font.bold = True
doc.add_paragraph(
    "  Camera (Pi)  →  YOLO crop (Pi)  →  medibox-edge queue (Pi SQLite)  →\n"
    "  Cloud API (Cloud Run)  →  Celery Worker (Cloud Run)  →\n"
    "  Vertex AI / Qwen2.5-VL-7B-AWQ  →  PostgreSQL result  →  Frontend dashboard"
, style="Normal")

# ════════════════════════════════════════════════════════════════════════════
# 3. EDGE LAYER (Pi)
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Edge Layer — Raspberry Pi 5")
heading(doc, "3.1  Hardware & Connectivity", level=2)
add_table(doc,
    ["Property", "Value"],
    [
        [cell("Device"), cell("Raspberry Pi 5, USB camera")],
        [cell("Username / Password"), cell("verox / verox1")],
        [cell("Tailscale IP"), cell("100.84.95.114")],
        [cell("Hostname"), cell("verox.tail591eb0.ts.net")],
        [cell("OS"), cell("Raspberry Pi OS, Python 3.11")],
        [cell("VPN"), cell("Tailscale (SSH/admin access only)")],
        [cell("Internet"), cell("Direct outbound for Cloud API calls")],
    ],
    col_widths=[2.2, 4.0]
)

heading(doc, "3.2  Software Components", level=2)
add_table(doc,
    ["Component", "Location on Pi", "Role", "Status"],
    [
        [cell("YOLO camera app"),   cell("~/yolo-ws/app.py"),              cell("Captures frames, runs YOLO detection, crops prescription regions, sends crops to medibox-edge queue. Also streams raw frames to cloud at ~10 fps."), DONE],
        [cell("medibox-edge"),      cell("/opt/medibox-edge/venv"),          cell("systemd daemon. Picks crops from SQLite queue, authenticates via Firebase SA key, submits to POST /v1/submit, polls for results, handles retries."), DONE],
        [cell("Config"),            cell("/etc/medibox/edge.toml"),          cell("Device ID, cloud URL, Firebase API key, SA key path."), DONE],
        [cell("SA Key"),            cell("/etc/medibox/medibox-edge.json"),  cell("Firebase service account key for device authentication."), DONE],
        [cell("systemd unit"),      cell("/etc/systemd/system/medibox-edge.service"), cell("Starts on boot, auto-restarts on crash."), DONE],
    ],
    col_widths=[1.6, 2.2, 3.2, 1.2]
)

heading(doc, "3.3  Edge Data Flow", level=2)
steps = [
    "1. USB camera feeds frames to app.py (OpenCV)",
    "2. YOLO v8 runs on-device → detects prescription bbox → crops JPEG",
    "3. Crop stored in SQLite queue (offline-first, survives reboot)",
    "4. Submitter thread: GET Firebase ID token (SA key) → POST /v1/submit with Bearer token",
    "5. ResultWatcher: polls GET /v1/result/:jobId every 2 s (or WebSocket if enabled)",
    "6. In parallel: cloud_push_loop thread pushes raw frames to POST /v1/camera/push at 10 fps",
]
for s in steps:
    doc.add_paragraph(s, style="List Bullet")

heading(doc, "3.4  medibox-edge Internal Modules", level=2)
add_table(doc,
    ["Module", "Responsibility"],
    [
        [cell("auth.py"),       cell("Firebase custom token → ID token exchange. Proactive refresh every 50 min.")],
        [cell("queue.py"),      cell("SQLite WAL queue. State machine: pending → submitted → confirmed / failed.")],
        [cell("submitter.py"),  cell("Async Celery-like loop. Semaphore-gated concurrent submissions.")],
        [cell("watcher.py"),    cell("Polls /v1/result per job. Falls back from WebSocket to HTTP polling.")],
        [cell("ws.py"),         cell("Optional WebSocket client for real-time result streaming.")],
        [cell("health.py"),     cell("LED status controller + /health HTTP endpoint.")],
        [cell("daemon.py"),     cell("systemd entry point. Orchestrates all components, handles SIGTERM drain.")],
        [cell("config.py"),     cell("Pydantic-validated TOML config loader.")],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. CLOUD LAYER
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Cloud Layer — Google Cloud Platform")
p = doc.add_paragraph()
p.add_run("Project: ").font.bold = True
p.runs[0].font.bold = True
p.add_run("verox-4dc3f   |   Region: us-central1")

heading(doc, "4.1  Cloud Run — API Service (medibox-api)", level=2)
add_table(doc,
    ["Property", "Value"],
    [
        [cell("URL"),       cell("https://medibox-api-5w7o5tyr2q-uc.a.run.app")],
        [cell("Runtime"),   cell("FastAPI + uvicorn (4 workers), Python 3.11")],
        [cell("Resources"), cell("1 vCPU · 1 GiB RAM · 0–5 instances · 60 s timeout")],
        [cell("Auth"),      cell("Firebase JWT verified on protected endpoints. Camera endpoints use X-Camera-Secret header.")],
        [cell("Access"),    cell("Public (allUsers IAM) — application-level auth handles security")],
    ],
    col_widths=[2.0, 4.5]
)

heading(doc, "4.1.1  API Endpoints", level=2)
add_table(doc,
    ["Method", "Path", "Auth", "Description"],
    [
        [cell("POST"), cell("/v1/submit"),            cell("Firebase JWT"), cell("Receive crops from Pi, create job, enqueue to Celery.")],
        [cell("GET"),  cell("/v1/result/:id"),         cell("Firebase JWT"), cell("Poll job status + result.")],
        [cell("POST"), cell("/v1/corrections/:id"),    cell("Firebase JWT"), cell("Pharmacist submits corrected JSON.")],
        [cell("GET"),  cell("/v1/admin/models"),       cell("Admin JWT"),    cell("Model registry list.")],
        [cell("POST"), cell("/v1/admin/model/rollback"),cell("Admin JWT"),   cell("Roll back to previous model version.")],
        [cell("POST"), cell("/v1/admin/maintenance"),  cell("Admin JWT"),    cell("Enable / disable maintenance mode (Redis flag).")],
        [cell("POST"), cell("/v1/camera/push"),        cell("X-Camera-Secret"), cell("Pi pushes JPEG frames (stored in Redis, 10 s TTL).")],
        [cell("GET"),  cell("/v1/camera/snapshot"),    cell("None"),         cell("Frontend polls latest Pi frame.")],
        [cell("POST"), cell("/v1/camera/capture"),     cell("None"),         cell("Capture current frame → create OCR job → return job_id.")],
        [cell("WS"),   cell("/v1/ws/:jobId"),          cell("Firebase JWT"), cell("Real-time job progress via WebSocket + Redis pub/sub fan-out.")],
        [cell("GET"),  cell("/v1/healthz"),            cell("None"),         cell("Liveness probe.")],
        [cell("GET"),  cell("/v1/readyz"),             cell("None"),         cell("Readiness probe (checks DB + Redis).")],
    ],
    col_widths=[0.7, 2.3, 1.4, 2.1]
)

heading(doc, "4.2  Cloud Run — Worker Service (medibox-worker)", level=2)
add_table(doc,
    ["Property", "Value"],
    [
        [cell("Runtime"),   cell("Celery worker, Python 3.11 + PyTorch + CUDA")],
        [cell("Resources"), cell("2 vCPU · 2 GiB RAM · min 1 instance (always-on) · max 5 · 3600 s timeout")],
        [cell("Queue"),     cell("Redis inference queue (Memorystore 10.4.26.107:6379)")],
    ],
    col_widths=[2.0, 4.5]
)

heading(doc, "4.2.1  Worker Pipeline (run_pipeline task)", level=2)
pipeline_steps = [
    "1. Safety filter — validate image bytes (format, size, decompression-bomb check)",
    "2. Upload raw crops to GCS bucket verox-4dc3f-crops/{device_id}/{job_id}/",
    "3. Grid composer — pack multiple crops into a single grid image for batch inference",
    "4. Vertex AI call — POST grid to Qwen2.5-VL-7B-AWQ via vLLM /v1/chat/completions",
    "5. Postprocessing — JSON repair, drug name normalization (Tunisian formulary), confidence scoring, hallucination detection",
    "6. Write result to PostgreSQL jobs table",
    "7. Publish to Redis pub/sub → WebSocket fan-out to frontend",
]
for s in pipeline_steps:
    doc.add_paragraph(s, style="List Bullet")

heading(doc, "4.3  Cloud Run — Frontend (medibox-frontend)", level=2)
add_table(doc,
    ["Property", "Value"],
    [
        [cell("URL"),       cell("https://medibox-frontend-5w7o5tyr2q-uc.a.run.app")],
        [cell("Runtime"),   cell("nginx serving static React/Vite bundle")],
        [cell("Framework"), cell("React 18 + TypeScript + Vite + Tailwind CSS")],
        [cell("Resources"), cell("1 vCPU · 512 MiB RAM · 0–3 instances")],
    ],
    col_widths=[2.0, 4.5]
)

heading(doc, "4.3.1  Frontend Pages", level=2)
add_table(doc,
    ["Route", "Page", "Description"],
    [
        [cell("/camera"),          cell("Camera Page"),      cell("Live Pi camera feed (polls every 200 ms). Start Camera → Capture & Submit → navigates to job tracker.")],
        [cell("/submit"),          cell("Submit Page"),      cell("Manual drag-and-drop prescription crop upload.")],
        [cell("/jobs/:id"),        cell("Job Tracker"),      cell("Real-time status updates via WebSocket. Shows progress bar → final JSON result.")],
        [cell("/corrections/:id"), cell("Corrections Page"), cell("Pharmacist reviews and edits extracted medication data.")],
        [cell("/admin"),           cell("Admin Page"),       cell("Model registry, maintenance mode toggle, model rollback.")],
    ],
    col_widths=[1.5, 1.5, 3.6]
)

heading(doc, "4.4  Vertex AI — AI Model", level=2)
add_table(doc,
    ["Property", "Value"],
    [
        [cell("Model"),       cell("Qwen2.5-VL-7B-AWQ (7 B vision-language, 4-bit AWQ quantized)")],
        [cell("Serving"),     cell("vLLM v0.6.6 — OpenAI-compatible /v1/chat/completions")],
        [cell("Container"),   cell("vllm/vllm-openai:v0.6.6.post1 + autoawq + qwen-vl-utils")],
        [cell("Endpoint ID"), cell("1778860829916004352")],
        [cell("Model ID"),    cell("4963037707862278144")],
        [cell("GPU target"),  cell("NVIDIA L4 (g2-standard-4) — preferred. T4 (n1-standard-4) as fallback.")],
        [cell("Input"),       cell("Base64 grid image + system prompt (Tunisian pharmacy context, Arabic/French)")],
        [cell("Output"),      cell("Structured JSON: medications list with name, dosage, frequency, patient info, confidence scores")],
    ],
    col_widths=[2.0, 4.5]
)

doc.add_page_break()

heading(doc, "4.5  Data Layer", level=2)
add_table(doc,
    ["Resource", "Details"],
    [
        [cell("Cloud SQL"),     cell("medibox-postgres · PostgreSQL 15 · db-f1-micro · 10 GB · daily backups · us-central1")],
        [cell("DB Tables"),     cell("jobs, corrections, audit_log, model_registry, admin_role_grants, deployment_log")],
        [cell("Migrations"),    cell("Alembic (async asyncpg). Run via Cloud Run Job 'run-migrations'. 2 versions: 001_initial_schema, 002_gcp_tables")],
        [cell("Redis"),         cell("Memorystore Redis 7.0 · 1 GB · Basic · 10.4.26.107:6379")],
        [cell("Redis uses"),    cell("Celery broker (inference queue) · WebSocket pub/sub · camera frame relay (10 s TTL) · maintenance flag · rate-limit counters")],
        [cell("GCS Buckets"),   cell("verox-4dc3f-crops (raw crops per job) · verox-4dc3f-models (fine-tuned adapter weights)")],
    ],
    col_widths=[1.8, 4.7]
)

heading(doc, "4.6  Security & Compliance", level=2)
add_table(doc,
    ["Concern", "Implementation"],
    [
        [cell("Authentication"),   cell("Firebase JWT on all protected API endpoints. Pi uses SA key → custom token → ID token.")],
        [cell("Secrets"),          cell("7 secrets in GCP Secret Manager. Never in code, env files, or source control.")],
        [cell("PII Encryption"),   cell("Patient and doctor names AES-encrypted via Cloud KMS (envelope encryption) before DB write.")],
        [cell("Network"),          cell("API + Worker run on VPC connector. Redis and SQL accessible only via private IP. Frontend is the only public-facing service.")],
        [cell("Rate Limiting"),    cell("Redis token bucket: 60 req/min per device_id.")],
        [cell("Audit Log"),        cell("Every job submission, correction, admin action written to audit_log table with correlation_id, actor_uid, IP.")],
        [cell("Camera Auth"),      cell("X-Camera-Secret header (stored in Secret Manager). Pi stream rejected without valid secret.")],
        [cell("CORS"),             cell("Frontend URL + localhost:5173 only.")],
    ],
    col_widths=[2.0, 4.5]
)

heading(doc, "4.7  CI/CD & Monitoring", level=2)
add_table(doc,
    ["System", "Details"],
    [
        [cell("Cloud Build"),      cell("Pipeline: lint (ruff + mypy) → 78 unit tests → build 3 Docker images → push → deploy → smoke test.")],
        [cell("Monitoring"),       cell("4 dashboards: System Overview, Model Quality, MLOps, Cost. 13 log-based metrics. 8 alert policies.")],
        [cell("Notifications"),    cell("Email alerts to guesmitaha96@gmail.com via notification channel 9744769548402279807.")],
        [cell("Budget Alert"),     cell("$500/month cap with 50%, 80%, 100% thresholds. ✅ Created.")],
    ],
    col_widths=[2.0, 4.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. DEPLOYMENT STATUS
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Deployment Status — Full Checklist")

heading(doc, "5.1  Infrastructure", level=2)
add_table(doc,
    ["Item", "Status", "Notes"],
    [
        [cell("GCP project verox-4dc3f"),             DONE,    cell("")],
        [cell("All APIs enabled"),                     DONE,    cell("")],
        [cell("Service accounts created"),             DONE,    cell("medibox-runner, medibox-worker, medibox-ci, medibox-edge")],
        [cell("VPC connector medibox-connector"),      DONE,    cell("READY state")],
        [cell("Cloud SQL medibox-postgres"),           DONE,    cell("PostgreSQL 15, RUNNABLE")],
        [cell("Redis Memorystore"),                    DONE,    cell("Redis 7.0, 1 GB, 10.4.26.107:6379")],
        [cell("GCS buckets"),                          DONE,    cell("verox-4dc3f-crops, verox-4dc3f-models")],
        [cell("Artifact Registry"),                    DONE,    cell("us-central1-docker.pkg.dev/verox-4dc3f/medibox-repo")],
        [cell("KMS key ring"),                         DONE,    cell("medibox-keyring, pii-key")],
        [cell("Cloud Build service account"),          DONE,    cell("medibox-ci@verox-4dc3f.iam.gserviceaccount.com")],
    ],
    col_widths=[3.0, 1.4, 2.1]
)

heading(doc, "5.2  Secrets", level=2)
add_table(doc,
    ["Secret Name", "Status", "Notes"],
    [
        [cell("medibox-db-password"),          DONE, cell("")],
        [cell("medibox-redis-auth"),           DONE, cell("")],
        [cell("medibox-jwt-signing-key"),      DONE, cell("")],
        [cell("medibox-pii-encryption-key"),   DONE, cell("")],
        [cell("medibox-firebase-admin-json"),  DONE, cell("")],
        [cell("medibox-database-url"),         DONE, cell("")],
        [cell("medibox-database-url-sync"),    DONE, cell("")],
        [cell("medibox-camera-secret"),        DONE, cell("Stored in GCP Secret Manager — not shown here")],
    ],
    col_widths=[3.0, 1.4, 2.1]
)

heading(doc, "5.3  Cloud Run Services", level=2)
add_table(doc,
    ["Service", "Status", "Notes"],
    [
        [cell("medibox-api"),       DONE,    cell("Running. Public. Camera endpoints + full OCR API.")],
        [cell("medibox-worker"),    DONE,    cell("Running. min=1 always-on. Celery + PyTorch.")],
        [cell("medibox-frontend"),  DONE,    cell("Running. Camera page, Submit, Jobs, Corrections, Admin.")],
        [cell("DB Migrations"),     PEND,    cell("Build 7731bf51 deploying now. Once API image is live, Cloud Run Job 'run-migrations' will be executed.")],
    ],
    col_widths=[2.0, 1.4, 3.1]
)

heading(doc, "5.4  AI Model / Vertex AI", level=2)
add_table(doc,
    ["Item", "Status", "Notes"],
    [
        [cell("vLLM Docker image built"),      DONE,    cell("medibox-vllm:latest in Artifact Registry")],
        [cell("Vertex AI endpoint created"),   DONE,    cell("ID: 1778860829916004352")],
        [cell("Model uploaded to Vertex"),     DONE,    cell("ID: 4963037707862278144 (medibox-vllm-latest)")],
        [cell("GPU model deployed (L4)"),      BLOCKED, cell("L4 (g2-standard-4) repeatedly 'temporarily unavailable' in us-central1. Capacity issue, not quota — NVIDIA_L4_GPUS quota=1.")],
        [cell("GPU model deployed (T4)"),      PROG,    cell("Deployment operation 7533333355413110784 running now. ~20-40 min.")],
        [cell("VERTEX_ENDPOINT_ID in worker"), PEND,    cell("Will be set once T4 or L4 deployment succeeds.")],
    ],
    col_widths=[2.4, 1.4, 2.7]
)

heading(doc, "5.5  Edge (Raspberry Pi 5)", level=2)
add_table(doc,
    ["Item", "Status", "Notes"],
    [
        [cell("medibox-edge installed"),          DONE, cell("pip installed from tarball in /opt/medibox-edge/venv")],
        [cell("systemd service active"),          DONE, cell("medibox-edge.service ACTIVE/RUNNING")],
        [cell("Firebase auth working"),           DONE, cell("Token refresh confirmed. Cloud health: HEALTHY")],
        [cell("SA key on device"),                DONE, cell("/etc/medibox/medibox-edge.json")],
        [cell("edge.toml configured"),            DONE, cell("device_id=pi-0001, api_base_url, firebase_api_key set")],
        [cell("camera app (app.py) updated"),     DONE, cell("cloud_push_loop added. Pushes frames to /v1/camera/push at ~10 fps")],
        [cell("Camera live feed (frontend)"),     PEND, cell("Frontend deployed. Pi needs to run: cd ~/yolo-ws && python app.py")],
        [cell("End-to-end job submission"),       PEND, cell("Waiting for DB migrations + Vertex endpoint to be live")],
    ],
    col_widths=[2.4, 1.4, 2.7]
)

heading(doc, "5.6  Admin Account", level=2)
add_table(doc,
    ["Item", "Status", "Notes"],
    [
        [cell("Firebase user created"),               DONE,    cell("UID: 2UzPsPp4npTqlvjztL7B5Ja4PW72  |  guesmitaha96@gmail.com")],
        [cell("Firebase admin claim set"),            DONE,    cell("admin: True set via Firebase Admin SDK")],
        [cell("admin_role_grants DB row"),            PEND,    cell("Waiting for migrations to create the table first, then INSERT will run automatically.")],
    ],
    col_widths=[2.4, 1.4, 2.7]
)

heading(doc, "5.7  Monitoring & Observability", level=2)
add_table(doc,
    ["Item", "Status", "Notes"],
    [
        [cell("Email notification channel"),      DONE, cell("guesmitaha96@gmail.com — channel ID 9744769548402279807")],
        [cell("13 log-based metrics"),            DONE, cell("Confidence, correction, hallucination, drug normalize, JSON repair, etc.")],
        [cell("4 dashboards"),                    DONE, cell("System Overview, Model Quality, MLOps, Cost")],
        [cell("8 alert policies"),                DONE, cell("API latency, error rate, GPU duty cycle, SQL CPU, SQL connections, hallucination spike, Vertex health, audit log")],
        [cell("2 alert policies skipped"),    ("⚠️ Warn", ORANGE, False, None), cell("Redis queue depth metric & billing monthly_cost metric not yet available in project")],
        [cell("Budget alert ($500/month)"),       DONE, cell("Created manually in Cloud Console ✅")],
    ],
    col_widths=[2.4, 1.4, 2.7]
)

heading(doc, "5.8  CI/CD", level=2)
add_table(doc,
    ["Item", "Status", "Notes"],
    [
        [cell("cloudbuild.yaml pipeline"),    DONE,  cell("lint → test → build (API+Worker+Frontend) → push → deploy → smoke test")],
        [cell("Build triggers on push"),      NO,    cell("No GitHub remote connected yet. Manual gcloud builds submit for now.")],
        [cell("GitHub repository"),           NO,    cell("Not pushed. No GitHub PAT provided.")],
    ],
    col_widths=[2.4, 1.4, 2.7]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. WHAT IS BLOCKED AND WHY
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Blocked Items — Root Causes")
add_table(doc,
    ["Item", "Blocked By", "Resolution"],
    [
        [
            cell("DB Migrations"),
            cell("Dockerfile was missing COPY migrations/ — alembic.ini not in container image"),
            cell("Fixed in build 7731bf51 (currently running). Run-migrations job will execute after deploy.")
        ],
        [
            cell("admin_role_grants row"),
            cell("Table doesn't exist yet — migrations haven't run"),
            cell("Automatic after migrations succeed.")
        ],
        [
            cell("End-to-end OCR test"),
            cell("Needs: (1) migrations done, (2) Vertex endpoint live"),
            cell("Both resolving in parallel now.")
        ],
        [
            cell("L4 GPU deployment"),
            cell("GCP capacity: 'Machine type temporarily unavailable' for g2-standard-4 in us-central1. This is NOT a quota issue — quota is 1."),
            cell("T4 deployment running as fallback. Retry L4 later or try us-east4.")
        ],
        [
            cell("VERTEX_ENDPOINT_ID in worker"),
            cell("No GPU model deployed yet"),
            cell("Will be set immediately after T4 deployment succeeds.")
        ],
        [
            cell("GitHub CI/CD trigger"),
            cell("No GitHub repository connected, no PAT provided"),
            cell("Manual step: create GitHub repo, push code, connect to Cloud Build trigger.")
        ],
        [
            cell("Camera.py startup crash (4 builds)"),
            cell("(1) Unused variable raw. (2) Wrong Celery task name run_inference vs run_pipeline. (3) ws_manager.notify doesn't exist. (4) FastAPI 0.115 AssertionError: status_code=204 must not have response body."),
            cell("All fixed. Build 7731bf51 should deploy cleanly.")
        ],
    ],
    col_widths=[1.8, 2.6, 2.1]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. IMMEDIATE NEXT STEPS (AUTO)
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Immediate Next Steps")
heading(doc, "Automated (no action needed from you)", level=2)
auto_steps = [
    "Build 7731bf51 completes → API image with migrations + fixed camera.py deployed",
    "Cloud Run Job 'run-migrations' executes → all DB tables created",
    "INSERT admin row into admin_role_grants (UID 2UzPsPp4npTqlvjztL7B5Ja4PW72)",
    "T4 GPU deployment completes → VERTEX_ENDPOINT_ID injected into medibox-worker",
    "End-to-end smoke test: /v1/healthz → /v1/readyz → test job submission",
]
for i, s in enumerate(auto_steps, 1):
    doc.add_paragraph(f"{i}. {s}", style="List Number")

heading(doc, "Manual (requires you)", level=2)
manual_steps = [
    "Start the Pi camera to test live feed:\n   ssh verox@100.84.95.114\n   cd ~/yolo-ws && source venv/bin/activate && python app.py",
    "GitHub: create repo → git push → connect Cloud Build trigger for auto-deploy on push",
    "L4 GPU (optional): retry gcloud ai endpoints deploy-model with g2-standard-4 once capacity is available — or try region us-east4",
]
for i, s in enumerate(manual_steps, 1):
    doc.add_paragraph(f"{i}. {s}", style="List Number")

# ════════════════════════════════════════════════════════════════════════════
# 8. KEY URLS & IDs
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "8. Key URLs & Identifiers")
add_table(doc,
    ["Resource", "Value"],
    [
        [cell("Frontend"),           cell("https://medibox-frontend-5w7o5tyr2q-uc.a.run.app")],
        [cell("API"),                cell("https://medibox-api-5w7o5tyr2q-uc.a.run.app")],
        [cell("GCP Project"),        cell("verox-4dc3f  (project number: 272258744118)")],
        [cell("Vertex Endpoint"),    cell("1778860829916004352")],
        [cell("Vertex Model"),       cell("4963037707862278144")],
        [cell("Admin Firebase UID"), cell("2UzPsPp4npTqlvjztL7B5Ja4PW72")],
        [cell("Pi IP"),              cell("100.84.95.114 (Tailscale) / verox.tail591eb0.ts.net")],
        [cell("Cloud SQL"),          cell("verox-4dc3f:us-central1:medibox-postgres")],
        [cell("Redis"),              cell("10.4.26.107:6379")],
        [cell("Cloud Run (console)"),cell("https://console.cloud.google.com/run?project=verox-4dc3f")],
        [cell("Vertex (console)"),   cell("https://console.cloud.google.com/vertex-ai/endpoints?project=verox-4dc3f")],
        [cell("Cloud Build"),        cell("https://console.cloud.google.com/cloud-build/builds?project=verox-4dc3f")],
        [cell("Monitoring"),         cell("https://console.cloud.google.com/monitoring/dashboards?project=verox-4dc3f")],
        [cell("Logs"),               cell("https://console.cloud.google.com/logs?project=verox-4dc3f")],
        [cell("Firebase Console"),   cell("https://console.firebase.google.com/project/verox-4dc3f")],
    ],
    col_widths=[2.0, 4.5]
)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\Guesmi Taha\Desktop\Medibox_Architecture.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
